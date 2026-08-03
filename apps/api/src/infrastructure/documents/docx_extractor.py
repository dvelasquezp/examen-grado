"""Extractor de DOCX con python-docx."""

from pathlib import Path

from docx import Document as DocxDocument

from src.domain.ingestion.chunk import ExtractedBlock, ExtractedDocument
from src.infrastructure.documents.normalizer import sanitize_text


class DocxExtractor:
    def extract(self, filepath: Path) -> ExtractedDocument:
        doc = DocxDocument(filepath)
        blocks: list[ExtractedBlock] = []

        for para in doc.paragraphs:
            text = sanitize_text(para.text.strip())
            if not text or len(text) < 2:
                continue

            block_type = "paragraph"
            heading_level = None

            if para.style and para.style.name.startswith("Heading"):
                block_type = "heading"
                try:
                    heading_level = int(para.style.name.replace("Heading", "").strip() or "1")
                except ValueError:
                    heading_level = 1
            elif para.runs and all(r.bold for r in para.runs if r.text.strip()):
                if len(text) < 200:
                    block_type = "heading"
                    heading_level = 2

            blocks.append(
                ExtractedBlock(
                    text=text,
                    page_number=None,
                    block_type=block_type,
                    heading_level=heading_level,
                )
            )

        return ExtractedDocument(blocks=blocks, page_count=0)
